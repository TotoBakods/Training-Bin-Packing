"""
parse_thesis.py
===============
Extracts Chapters 2 (RRL), 3 (Methodology), and 4 (Results/Data) from the
thesis .docx file and outputs a structured markdown document to:
  Documents/04_Machine_Learning/Performance_Metrics/thesis_chapter_requirements_and_rrl.md

Usage:
    python parse_thesis.py
"""

import os
from docx import Document

DOCX_PATH = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "OPTIMIZING WAREHOUSE STORAGE ALLOCATION USING GENETIC ALGORITHM AND "
    "EXTREMAL OPTIMIZATION FOR EFFICIENT SPACE  UTLIZTION AND INVENTORY "
    "MANAGEMENT (FINAL) (1).docx"
)

OUTPUT_PATH = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "Documents", "04_Machine_Learning", "Performance_Metrics",
    "thesis_chapter_requirements_and_rrl.md"
)

# Chapter boundary markers — keys are match strings, values are chapter keys
# None as value = stop-marker (chapter 5+, stop collecting)
CHAPTER_MARKERS = {
    "CHAPTER 2": "chapter_2",
    "CHAPTER II": "chapter_2",
    "CHAPTER 3": "chapter_3",
    "CHAPTER III": "chapter_3",
    "CHAPTER 4": "chapter_4",
    "CHAPTER IV": "chapter_4",
    "CHAPTER 5": None,
    "CHAPTER V": None,
}

TARGET_CHAPTERS = {"chapter_2", "chapter_3", "chapter_4"}

SECTION_TITLES = {
    "chapter_2": "## Chapter 2: Review of Related Literature (RRL)",
    "chapter_3": "## Chapter 3: Methodology",
    "chapter_4": "## Chapter 4: Results and Discussion",
}


def detect_chapter(paragraph):
    """
    Return the chapter key if this paragraph marks a chapter boundary.
    Returns False if not a chapter boundary at all.
    Returns None if it's a stop-marker (Chapter 5+).

    Matches longest markers first to avoid "CHAPTER II" matching "CHAPTER III".
    Requires the marker to be followed by a word boundary (space, colon, or EOL).
    """
    text_upper = paragraph.text.strip().upper()

    def matches(text, marker):
        """True only if marker occupies a whole word at the start of text."""
        if not text.startswith(marker):
            return False
        rest = text[len(marker):]
        return rest == "" or rest[0] in (" ", ":", "\n", "\t", "-")

    # Sort by marker length descending so "CHAPTER III" beats "CHAPTER II"
    sorted_markers = sorted(CHAPTER_MARKERS.items(), key=lambda kv: -len(kv[0]))

    # Primary check: text content (works regardless of style)
    for marker, key in sorted_markers:
        if matches(text_upper, marker):
            return key

    # Fallback: Heading 1 style containing chapter keyword anywhere
    if "HEADING 1" in paragraph.style.name.upper():
        for marker, key in sorted_markers:
            if marker in text_upper:
                return key

    return False  # not a chapter boundary


def is_heading(paragraph):
    return "HEADING" in paragraph.style.name.upper()


def extract_chapters(doc):
    """Walk paragraphs and collect text grouped by target chapter."""
    chapters = {k: [] for k in TARGET_CHAPTERS}
    current_chapter = False  # False = not yet in a target chapter

    for para in doc.paragraphs:
        boundary = detect_chapter(para)

        if boundary is not False:
            # Hit a chapter boundary
            current_chapter = boundary  # may be a key or None (stop-marker)
            continue

        if current_chapter and current_chapter in TARGET_CHAPTERS:
            text = para.text.strip()
            if text:
                prefix = "### " if is_heading(para) else ""
                chapters[current_chapter].append(prefix + text)

    return chapters


def table_to_markdown(table):
    """Convert a python-docx Table to a pipe-delimited markdown string."""
    rows = []
    for i, row in enumerate(table.rows):
        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("|" + "|".join([" --- " for _ in cells]) + "|")
    return "\n".join(rows)


def extract_chapter4_tables(doc):
    """
    python-docx paragraphs skips table cells, so we collect all tables
    separately and append them to chapter_4.
    Tables are appended at the end since we cannot easily determine their
    chapter position relative to paragraphs without the full XML tree walk.
    """
    table_lines = []
    for table in doc.tables:
        md = table_to_markdown(table)
        if md.strip():
            table_lines.append(md)
    return table_lines


def build_markdown(chapters, table_lines):
    lines = [
        "# Thesis Chapter Requirements and RRL Extraction",
        "",
        "> Auto-extracted from thesis `.docx` using `parse_thesis.py`  ",
        "> Source: *OPTIMIZING WAREHOUSE STORAGE ALLOCATION USING GENETIC ALGORITHM",
        "> AND EXTREMAL OPTIMIZATION FOR EFFICIENT SPACE UTILIZATION AND INVENTORY MANAGEMENT*",
        "",
        "---",
        "",
    ]

    for key in ["chapter_2", "chapter_3", "chapter_4"]:
        lines.append(SECTION_TITLES[key])
        lines.append("")
        content = chapters.get(key, [])
        if content:
            lines.extend(content)
        else:
            lines.append(
                "> *(No content extracted — chapter heading format may differ from "
                "expected markers. Check `CHAPTER_MARKERS` in `parse_thesis.py`.)*"
            )
        lines.append("")

        # Append tables to Chapter 4 only
        if key == "chapter_4" and table_lines:
            lines.append("### Extracted Tables (from thesis .docx)\n")
            for t in table_lines:
                lines.append(t)
                lines.append("")

        lines.append("---")
        lines.append("")

    return "\n".join(lines)


def main():
    if not os.path.exists(DOCX_PATH):
        print(f"ERROR: Thesis file not found at:\n  {DOCX_PATH}")
        print("Ensure the .docx file is in the repository root directory.")
        return

    print(f"Opening: {os.path.basename(DOCX_PATH)}")
    doc = Document(DOCX_PATH)
    print(f"Total paragraphs: {len(doc.paragraphs)}")

    chapters = extract_chapters(doc)
    for key in TARGET_CHAPTERS:
        print(f"  {key}: {len(chapters[key])} paragraphs extracted")

    table_lines = extract_chapter4_tables(doc)
    print(f"  Tables found in document: {len(table_lines)}")

    md = build_markdown(chapters, table_lines)

    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    with open(OUTPUT_PATH, "w", encoding="utf-8") as f:
        f.write(md)

    print(f"\nOutput written to:\n  {OUTPUT_PATH}")
    print(f"Total lines: {len(md.splitlines())}")

    # Warn if any chapter came up empty
    for key in TARGET_CHAPTERS:
        if not chapters[key]:
            print(
                f"\nWARNING: {key} is empty. Open the .docx in Word and check the "
                f"exact text/style of the chapter heading, then update CHAPTER_MARKERS."
            )


if __name__ == "__main__":
    main()
